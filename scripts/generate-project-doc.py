import datetime
import json
import os
import re
import subprocess
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_PATH = os.path.expanduser('~/Desktop/Lofty Support Portal - Project Documentation.docx')

INDIGO = RGBColor(0x33, 0x2E, 0xB8)
INK = RGBColor(0x1C, 0x1F, 0x2E)
GRAY = RGBColor(0x5B, 0x62, 0x74)
GREEN = RGBColor(0x0F, 0x7A, 0x49)

def load_changelog():
    """Parses the live CHANGELOG array straight out of pto-public.html so this doc can
    never drift from what's actually shipped - no separately-maintained copy to forget."""
    html = open(os.path.join(REPO_DIR, 'pto-public.html')).read()
    m = re.search(r'const CHANGELOG=(\[.*?\n\]);', html, re.DOTALL)
    node_script = "const CHANGELOG=" + m.group(1) + ";console.log(JSON.stringify(CHANGELOG));"
    out = subprocess.run(['node', '-e', node_script], capture_output=True, text=True, check=True)
    return json.loads(out.stdout)

def load_portal_version():
    server = open(os.path.join(REPO_DIR, 'pto-public-server.js')).read()
    m = re.search(r"const PORTAL_VERSION = '([\d.]+)';", server)
    return m.group(1)

CHANGELOG = load_changelog()
PORTAL_VERSION = load_portal_version()

doc = Document()

# Base style
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = INK

def set_cell_shading(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def h1(text, color=INDIGO, size=22, space_before=18, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

def h2(text, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = INK
    return p

def body(text, size=11, color=INK, space_after=6, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return p

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r2 = p.add_run(text)
    else:
        r = p.add_run(text)
    return p

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'DFE2EA')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ---------- Title Page ----------
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
title_p.paragraph_format.space_before = Pt(60)
r = title_p.add_run('Lofty Support Portal')
r.bold = True
r.font.size = Pt(34)
r.font.color.rgb = INDIGO

sub_p = doc.add_paragraph()
sub_p.paragraph_format.space_after = Pt(4)
r = sub_p.add_run('Project Documentation')
r.font.size = Pt(18)
r.font.color.rgb = GRAY

meta_p = doc.add_paragraph()
meta_p.paragraph_format.space_before = Pt(10)
r = meta_p.add_run(f'Internal reference · Current version {PORTAL_VERSION} · Prepared by Mac · {datetime.date.today().strftime("%B %Y")}')
r.font.size = Pt(11)
r.font.color.rgb = GRAY
r.italic = True

hr()

# ---------- Overview ----------
h1('Overview')
body(
    'The Lofty Support Portal (repo name "mtd-kpi", public service name "loftysupport.onrender.com") is a '
    'role-based web application that replaced a scattered mix of spreadsheets, chat threads, and manual tracking '
    'for the support floor. Every employee, from a frontline rep to the Senior Operations Manager, signs into the '
    'same portal, and the interface adapts automatically based on who they are.'
)
body(
    'It covers day-to-day performance visibility (KPI, attendance, schedule), people-management workflows '
    '(coaching, disciplinary action, SOP alignment and sign-off), quality assurance (company-wide DSAT review), '
    'onboarding (structured new-hire tracking), and team culture (a live recognition wall and automated shift-end '
    'thank-yous).'
)

h2('Architecture at a glance')
body('The system runs as a two-server split:', space_after=4)
bullet('Local admin server (zendesk-proxy.js) - pulls live data from Zendesk/Jira, computes KPI, and is the source of truth for roster/attendance/schedule files.')
bullet('Public portal server (pto-public-server.js, deployed on Render as loftysupport.onrender.com) - the internet-facing app every employee logs into. It never talks to Zendesk directly.')
bullet('The two servers stay in sync through a request-queue in Upstash (Redis): the public server writes pending changes to a queue key, and the local admin server drains that queue on a ~30 second timer and applies it to the canonical local data files.')
bullet('The front end (pto-public.html) is a single-page app - no build step, no framework - talking to the public server over a small REST API defined in shared/*.js service modules.')

# ---------- Who uses it ----------
h1('Roles')
roles = [
    ('Rep', 'Sees their own KPI, schedule, coaching history, disciplinary history, PTO requests, and Alignment items assigned to them.'),
    ('Team Lead', "Everything a Rep sees, plus manages their direct reports' coaching, disciplinary filings, team attendance entry, team roster, and can author Alignment (SOP) rollouts."),
    ('BQA (Quality)', 'Company-wide DSAT review - flags bad CSAT ratings as Not Valid, approves/rejects rep-filed disputes, and can open a coaching log on a team lead\'s behalf once a DSAT stands as valid.'),
    ('SOM / HR', 'Final approval on disciplinary tiers and Alignment rollouts, plus a company-wide Alignment compliance summary (by site, team, and category).'),
    ('Training Manager', 'A scoped version of the Team Lead tools restricted to trainees only (kpiType = "Trainee") - onboarding pipeline, scorecards, and endorsement into a real production role and team.'),
    ('Admin', 'Full "View As" preview of any role for support/testing purposes; view-as only ever affects reads, never writes.'),
]
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Light Grid Accent 1'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr = tbl.rows[0].cells
hdr[0].text = 'Role'
hdr[1].text = 'What they can do'
for cell in hdr:
    set_cell_shading(cell, '332EB8')
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.bold = True
for name, desc in roles:
    row = tbl.add_row().cells
    row[0].text = name
    row[0].paragraphs[0].runs[0].bold = True
    row[1].text = desc
doc.add_paragraph()

# ---------- Feature areas ----------
h1('Feature Areas')

features = [
    ('Performance / KPI', [
        'Every rep sees their own Final KPI, category breakdown, and current period live, with company-wide Site KPI (call completion, long-call rate, CSAT) visible to anyone signed in.',
        'A Daily EOD Report breaks the same metrics down day by day, and the Site KPI tab shows who is currently in training and how long they have been onboarding.',
        'The internal admin dashboard exports a formatted "KPI Bonus Review Workbook" (.xlsx) - Final KPI is highlighted green and shown both as a quick-glance column up front and in full detail at the end of the row, the individual weighted scores that sum into it (CSAT/Calls/LCR/Ticket/FRT/Team/Lead Import/Bonus Points) are highlighted gold, and each KPI type/team lead block is visually separated so nothing runs together.',
        'Refreshing metrics for one period (e.g. this month) runs independently of whatever period is currently on screen, so a team lead can keep browsing last period\'s cached numbers while a refresh completes in the background, then switch over once it\'s ready.',
    ]),
    ('Coaching', [
        'A new coaching session can be explicitly linked to an earlier one covering the same underlying issue, and the UI shows an "Occurrence #N" badge computed from that chain - so a repeat issue is visibly a repeat, not a fresh first offense.',
        'Employees fill in the Discussion & Development Plan, Action Plan, and Follow-Up Date themselves when they review and sign, rather than the team lead pre-filling it.',
        'Every session gets a human-readable code (e.g. COACH-2026-0041) for reference.',
    ]),
    ('Disciplinary', [
        'A fixed path from infraction to sign-off: Team Lead files an Incident Report, SOM pre-reviews the tier, HR makes the final call, and the employee e-signs. Every step is timestamped and attributable.',
        'The system counts prior instances at the same tier automatically and suggests the appropriate escalation.',
    ]),
    ('Alignment (SOP rollout & sign-off)', [
        'A Team Lead publishes a policy or process update, picks who needs to see it (including leadership when relevant), optionally attaches an AI-generated comprehension quiz, and sets a due date; the SOM approves once.',
        'Each targeted employee reads the content, passes the quiz (questions and answer order are shuffled per rep to discourage copying), and e-signs by typing their name.',
        'As of v1.13.0, any Alignment item still awaiting review/quiz/signature opens automatically as an undismissable modal on login (and is re-checked every few minutes while signed in) - the rest of the portal is blocked until it is completed, so a rollout can no longer be missed because a tab never got clicked.',
        'A company-wide Alignment summary (by site, by team, by category) is available to the SOM.',
        'AI-assisted authoring ("Rephrase" and quiz generation) tries the company\'s internal Copilot service first and automatically falls back to Groq if Copilot is not connected or the call fails (v1.19.0) - same for the Coaching and Disciplinary narrative fields, which also use Rephrase.',
    ]),
    ('Onboarding / New Hires', [
        'The Training Manager adds a new hire directly into the roster and logs a scorecard as they progress (module completion, attendance, assessments).',
        'Endorsing a trainee lets the Training Manager pick their real production role and destination team lead in one step - Coaching, Disciplinary, Alignment, Team Attendance, and Team Roster access for that person then works automatically the same way it does for any other team lead\'s report.',
    ]),
    ('Quality / DSAT Review', [
        'BQA reviews every bad CSAT rating company-wide in one place (not scoped to a single team), and can flag a rating Not Valid or decide a rep-filed dispute in the same screen.',
        'Once a DSAT stands as valid, BQA can start a coaching log directly on the rep\'s real team lead\'s behalf, closing the loop between a bad customer experience and an actual coaching conversation.',
        'Ticket numbers in the review queue link straight out to the Zendesk ticket (added v1.15.0).',
        'AI-assisted triage (v1.16.0, network fix in v1.17.0): once the company\'s internal Copilot service is connected (a shared, admin-managed sign-in), every open ticket gets an automatic sentiment/risk read (High/Medium/Low) with a one-line reason, so BQA can prioritize the queue at a glance instead of opening each ticket cold. Copilot is only reachable from Lofty\'s own network, so the connect flow and the AI calls happen in the browser of whoever is using it, not on the server.',
    ]),
    ('Team Attendance', [
        'Team leads enter attendance for direct reports, which flows into KPI on the next refresh.',
        'A "Late" entry captures minutes tardy, an optional reason, and - as of v1.14.0 - whether the person was Onsite or WFH.',
    ]),
    ('Announcements', [
        'Team Leads, HR, and the SOM can post, edit, and delete company announcements from the portal.',
        'New announcements pop up automatically on login (and periodically while signed in) in a dismissible "Got it" modal; long announcements collapse behind a "Show more" toggle so one long post does not dominate the page (v1.15.0).',
        'Announcements can include an optional photo shown at the top of the post (v1.21.0).',
    ]),
    ('Spotlight Wall', [
        'A live, rotating recognition display: real customer shoutouts the moment a good rating comes in, plus shift-end "thank you" messages gated on genuine Zendesk/Jira activity so only people who actually worked get thanked.',
        'A daily broadcast thanks the whole team, agents and leaders alike.',
        'Occasional full-bleed poster slides (e.g. a CSAT/team-culture banner) and embedded Loom videos, sized to fill the display rather than sit in a small box (v2026-08-08).',
    ]),
    ('Rewards / Points (v1.20.0)', [
        'Points are computed on demand from existing data, never a hand-maintained ledger: good CSAT ratings, KPI performance tier, a clean attendance month, and a first-try Alignment quiz pass.',
        'A "Rewards Shop" lets reps redeem their balance for catalog items an SOM/HR/admin manages; redemptions go through a request -> approve/reject flow mirroring PTO, with only the approval trail tracked (not fulfillment).',
        'A monthly leaderboard ranks the team by points earned that month.',
    ]),
    ('PTO', [
        'Rep-facing filing and status tracking, team-lead pre-approval workflow, and forecast/threshold logic to flag risky requests before they are approved.',
    ]),
]
for title, items in features:
    h2(title)
    for it in items:
        bullet(it)

doc.add_page_break()

# ---------- Version history ----------
h1('Version History')
body('In-app changelog entries, oldest to newest. Versions before 1.2.0 predate the in-app changelog and are not itemized here; see the Git history for that earlier period.', italic=True, color=GRAY, size=10)
for entry in CHANGELOG:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"v{entry['version']}")
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = INDIGO
    for item in entry['items']:
        bullet(item)

doc.add_page_break()

# ---------- Notable fixes ----------
h1('Notable Fixes Along the Way')
body(
    'A sample of production issues found and fixed during development - included because they shaped some of the '
    'app\'s current safeguards.',
    italic=True, color=GRAY, size=10
)
fixes = [
    ('Critical login outage', 'A newly added shared JS module was missing from the server\'s static-file allow-list, so it 401\'d, which broke the page\'s module script entirely - nobody could sign in. Fixed and the allow-list check is now something to verify with every new shared/*.js file.'),
    ('Login form rendering over stale content', 'Signing out did not hide every tab\'s content section (only an old, incomplete hardcoded list), so the login form could appear floating over a still-visible previous tab. Fixed by aligning the hide-list with setTab()\'s complete view list.'),
    ('Training Manager seeing the wrong roster', 'Team Roster (and by extension Coaching/Disciplinary/Team Attendance) used the same "my direct reports" logic as a normal team lead, so under an admin\'s View-As preview it showed the admin\'s own real reports instead of the Training Manager\'s actual trainees. Fixed with a dedicated scoping helper restricted to kpiType === "Trainee".'),
    ('Alignment due date locked after approval', 'The edit lock meant to protect reviewed content also blocked fixing a due-date typo once an item was approved. Fixed with a narrow endpoint that can edit just the due date regardless of status.'),
    ('Stale hardcoded footer version', 'The footer showed a literal "V1.2.0" string disconnected from the real version variable. Fixed by making the footer render dynamically from the live portal version.'),
    ('Shift-end "thank you" messages mostly not firing', 'The admin dashboard checks each rep once, in a narrow 3-7-minutes-before-shift-end window, for a live activity signal - but a single miss (e.g. Zendesk status already idle a beat early) permanently disqualified that rep for the rest of the day instead of retrying on the next check. A separate mismatch also meant the "recent activity" fallback signal was only ever populated for the last ~5 minutes, not the intended 15, making it fail far more than expected. Fixed by retrying every tick until the window actually closes, and by fetching activity signals with a lookback that matches the presence check.'),
    ('Attendance-driven KPI fields silently going blank', 'Three separate spots read attendance by looking up one exact "month|endDate" snapshot key instead of merging the full chain of snapshots for that month: the local admin\'s attendance-editing grid, the attendance-summary calculation that KPI scoring depends on, and the summary-reading endpoint the KPI dashboard uses. Since a save under a new endDate only contains the dates that specific batch touched (not a full carry-forward copy), any of the three could make eligibleWorkdays - and therefore Productivity and Final KPI - look "Missing" for real, already-entered attendance. All three fixed to merge the full chain, matching the one place that already did this correctly.'),
    ('Refresh Period vs. Display Period on the internal KPI dashboard', 'Refreshing metrics for one period (e.g. this month) used to force-switch the whole dashboard to that period, so there was no way to keep viewing a different period\'s cached numbers while a refresh ran in the background - and the refresh and display views shared one busy-state flag, so clicking Load during a refresh silently did nothing at all. Fixed with independent period fields and an independent busy-state for refreshing vs. viewing.'),
    ('Database Agent zero-survey CSAT scored one tier too low', 'The documented "neutral" score for a period with zero CSAT surveys is 32 points (the middle of 5 tiers), but the code arrived at that by scoring a flat 80% rate through the normal tier lookup - which actually lands in the tier below (28 points), not the middle one. Fixed to return the intended tier directly.'),
    ('KPI Methodology export sheet read as one long undifferentiated table', 'The exported workbook\'s methodology reference tab listed every KPI type\'s scoring rules in a single flat table with a repeated "Section" column, making it easy to mistake one type\'s rule for another\'s. Restructured into one title-banded block per KPI type (General/Voice/Non-Voice/Senior TSR/Database Agent/Attendance/CSAT Disputes), each with its own header row, matching the rest of the workbook\'s per-type sheet style.'),
]
for name, desc in fixes:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(name)
    r.bold = True
    r.font.color.rgb = INK
    body(desc, size=10.5, color=GRAY, space_after=8)

# ---------- Roadmap ----------
h1('Roadmap / Not Yet Built')
bullet('Trainee-specific KPI framework, once scoring criteria are finalized.')
bullet('Wider Customer Shoutout detection on the Spotlight Wall - email/chat praise, not just CSAT surveys.')
bullet('Deeper adherence and attendance analytics.')

doc.save(OUT_PATH)
print('saved', OUT_PATH)
